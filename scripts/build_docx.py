# -*- coding: utf-8 -*-
"""
Генератор .docx из .md документов проекта.

Содержимое документов живёт в .md (единый источник). Этот скрипт собирает из них
готовые к печати/подписи .docx в _build/docx/, повторяя структуру папок.

Запуск:  python scripts/build_docx.py
Зависимость: python-docx  (pip install python-docx)

Поддерживаемый markdown:
  # ## ###            заголовки
  !! текст            крупный центрированный текст титульного листа (вне оглавления)
  ! текст             центрированный текст титульного листа (вне оглавления)
  >> текст            абзац с выравниванием вправо (УТВЕРЖДАЮ, блок подписи)
  - текст             маркированный список
  > текст             цитата (курсив)
  | a | b |           таблица (строка-разделитель |---| отделяет шапку)
  **жирный**          жирный текст внутри строки
  ==текст==           выделение жёлтым маркером (новое в текущей редакции)
  ![подпись](файл.png)  изображение во всю ширину полосы с подписью под ним
  <!-- landscape -->  альбомная ориентация страницы (для широких таблиц)
  <!-- style:rich --> оформление: цветные заголовки, заливка таблиц, «Стр. X из Y»
  <!-- pagebreak -->  разрыв страницы
  <!-- vspace:N -->   N пустых строк (вертикальный отступ на титульном листе)
  <!-- section:landscape --> новый раздел с альбомной ориентацией (для чертежей)
  <!-- toc -->        автособираемое оглавление Word (кликабельное, с номерами страниц)
  <!-- ... -->        комментарий (игнорируется)
  ---                 разделитель (игнорируется)
  строки "1. ..."     остаются как есть (нумерация пунктов документа)
"""
import os
import re
import sys
from urllib.parse import unquote

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("Не установлен python-docx. Выполните: python -m pip install python-docx")
    sys.exit(1)

# Палитра оформления (совпадает с прежней редакцией Регламента)
CLR_H1 = RGBColor(0x1F, 0x38, 0x64)   # тёмно-синий — разделы
CLR_H2 = RGBColor(0x2E, 0x74, 0xB5)   # насыщенный голубой — подразделы
FILL_HEAD = "D9E2F3"                  # шапка таблиц — светло-голубой
FILL_ROW = "F2F2F2"                   # чётные строки таблиц — светло-серый

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_ROOT = os.path.join(ROOT, "_build", "docx")

# Какие папки конвертировать (документы для печати). Служебные файлы пропускаем.
SRC_DIRS = [
    "01_Кадровые_документы",
    "02_Положение_об_отделе",
    "03_Внедрение_и_коллектив",
    "04_Шаблоны_управления_отделом",
    "05_Регламент_конструкторов",
]
# README — служебный; Регламент_КМД_текст.md — выжимка из старого .docx, не собирается
SKIP_FILES = {"README.md", "Регламент_КМД_текст.md"}

FONT = "Times New Roman"


def set_base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    p = st.paragraph_format
    p.space_after = Pt(4)
    p.line_spacing = 1.15


def add_runs_with_bold(paragraph, text):
    """Разбор **жирного** и ==выделенного маркером== внутри строки."""
    # сначала делим по ==выделению==, затем каждую часть — по **жирному**
    for k, chunk in enumerate(re.split(r"==", text)):
        if chunk == "":
            continue
        marked = k % 2 == 1
        for i, part in enumerate(re.split(r"\*\*", chunk)):
            if part == "":
                continue
            run = paragraph.add_run(part)
            run.font.name = FONT
            if i % 2 == 1:  # части между ** — жирные
                run.bold = True
            if marked:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def shade(cell, fill):
    """Заливка ячейки таблицы."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def add_field(paragraph, instr):
    """Вставка поля Word (PAGE, NUMPAGES, TOC …)."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "…"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr_el, sep, txt, end):
        run._r.append(el)
    return run


def add_toc(doc):
    """Автособираемое оглавление по заголовкам 1–2 уровня."""
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-2" \\h \\z \\u')


def add_page_footer(section):
    """Колонтитул «Стр. X из Y»."""
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Стр. ")
    r.font.name = FONT
    r.font.size = Pt(10)
    add_field(p, "PAGE")
    r2 = p.add_run(" из ")
    r2.font.name = FONT
    r2.font.size = Pt(10)
    add_field(p, "NUMPAGES")
    for run in p.runs:
        run.font.name = FONT
        run.font.size = Pt(10)


def enable_field_update(doc):
    """Word обновит оглавление и номера страниц при открытии файла."""
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    doc.settings.element.append(el)


def start_landscape_section(doc):
    """Новый раздел документа с альбомной ориентацией — для широких чертежей."""
    s = doc.add_section(WD_SECTION.NEW_PAGE)
    s.page_width = Cm(21.0)
    s.page_height = Cm(29.7)
    if s.page_width < s.page_height:
        s.orientation = WD_ORIENT.LANDSCAPE
        s.page_width, s.page_height = s.page_height, s.page_width
    s.top_margin = Cm(1.5)
    s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(1.5)
    s.right_margin = Cm(1.5)
    return s


def add_image(doc, path, caption, base_dir):
    """Изображение во всю ширину полосы набора с подписью под ним."""
    path = unquote(path)  # пути в markdown могут быть с %20 вместо пробелов
    full = path if os.path.isabs(path) else os.path.join(base_dir, path)
    if not os.path.exists(full):
        print("    ! не найдено изображение:", path)
        return
    s = doc.sections[-1]
    usable_w = s.page_width - s.left_margin - s.right_margin
    # запас по высоте на заголовок над иллюстрацией и подпись под ней
    usable_h = s.page_height - s.top_margin - s.bottom_margin - Cm(2.4)
    pic = doc.add_picture(full, width=usable_w)
    if pic.height > usable_h:
        k = usable_h / pic.height
        pic.height = int(pic.height * k)
        pic.width = int(pic.width * k)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.italic = True
        run.font.name = FONT
        run.font.size = Pt(10)


def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def is_separator_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return all(re.fullmatch(r":?-{2,}:?", c or "-") for c in cells)


def parse_cells(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_table(doc, rows, rich=False):
    sep_idx = next((i for i, r in enumerate(rows) if is_separator_row(r)), None)
    header = parse_cells(rows[0]) if sep_idx else None
    body_start = sep_idx + 1 if sep_idx is not None else 0
    body = [parse_cells(r) for r in rows[body_start:]]
    ncols = max([len(header)] if header else [0] + [len(r) for r in body])
    if ncols == 0:
        return
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def fill(cells_text, bold=False, band=None):
        cells = table.add_row().cells
        for j in range(ncols):
            txt = cells_text[j] if j < len(cells_text) else ""
            para = cells[j].paragraphs[0]
            add_runs_with_bold(para, txt)
            for r in para.runs:
                r.font.size = Pt(11)
                if bold:
                    r.bold = True
            if rich and band:
                shade(cells[j], band)

    if header:
        fill(header, bold=True, band=FILL_HEAD)
        # шапка повторяется на каждой странице при разрыве таблицы
        tr = table.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        el = OxmlElement("w:tblHeader")
        el.set(qn("w:val"), "true")
        trPr.append(el)
    for i, row in enumerate(body):
        fill(row, band=(FILL_ROW if i % 2 == 0 else None))
    doc.add_paragraph()


def render(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    directives = {l.strip() for l in lines if l.strip().startswith("<!--")}
    landscape = "<!-- landscape -->" in directives
    rich = "<!-- style:rich -->" in directives

    doc = Document()
    set_base_style(doc)
    for section in doc.sections:
        # формат A4 (шаблон python-docx по умолчанию — Letter)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        if landscape and section.page_width < section.page_height:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = (
                section.page_height, section.page_width)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
        if rich:
            add_page_footer(section)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Директивы
        if stripped == "<!-- pagebreak -->":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            i += 1
            continue
        if stripped == "<!-- toc -->":
            add_toc(doc)
            i += 1
            continue
        if stripped == "<!-- section:landscape -->":
            start_landscape_section(doc)
            i += 1
            continue
        m_vs = re.fullmatch(r"<!-- vspace:(\d+) -->", stripped)
        if m_vs:
            for _ in range(int(m_vs.group(1))):
                doc.add_paragraph()
            i += 1
            continue

        # Изображение: ![подпись](путь)
        m = re.fullmatch(r"!\[(.*?)\]\((.+?)\)", stripped)
        if m:
            add_image(doc, m.group(2), m.group(1), os.path.dirname(os.path.abspath(md_path)))
            i += 1
            continue

        if not stripped or stripped == "---" or stripped.startswith("<!--"):
            i += 1
            continue

        # Таблица
        if is_table_row(line):
            block = []
            while i < len(lines) and is_table_row(lines[i]):
                block.append(lines[i])
                i += 1
            add_table(doc, block, rich=rich)
            continue

        # Титульный лист: крупный и обычный центрированный текст (вне оглавления)
        if stripped.startswith("!! ") or stripped.startswith("! "):
            big = stripped.startswith("!! ")
            text = stripped[3:].strip() if big else stripped[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.name = FONT
            run.font.size = Pt(20 if big else 14)
            if rich and big:
                run.font.color.rgb = CLR_H1
            i += 1
            continue

        # Выравнивание вправо
        if stripped.startswith(">>"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_runs_with_bold(p, stripped[2:].strip())
            i += 1
            continue

        # Заголовки. В rich-режиме используются стили Heading 1/2/3 —
        # они попадают в автособираемое оглавление и подсвечиваются цветом.
        m_h = re.match(r"(#{1,3}) +(.*)", stripped)
        if m_h:
            level = len(m_h.group(1))
            style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
            size = {1: 15 if rich else 14, 2: 13, 3: 12}[level]
            color = CLR_H1 if level == 1 else CLR_H2
            p = doc.add_paragraph(style=style if rich else None)
            if level == 1 and not rich:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # через add_runs_with_bold — чтобы работало ==выделение== в заголовке
            add_runs_with_bold(p, m_h.group(2).strip())
            for run in p.runs:
                run.bold = True
                run.font.name = FONT
                run.font.size = Pt(size)
                if rich:
                    run.font.color.rgb = color
            i += 1
            continue

        # Цитата
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            add_runs_with_bold(p, stripped.lstrip(">").strip())
            for run in p.runs:
                run.italic = True
                run.font.name = FONT
            i += 1
            continue

        # Маркированный список
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, stripped[2:].strip())
            i += 1
            continue

        # Обычный абзац (включая пункты "1. ...")
        p = doc.add_paragraph()
        add_runs_with_bold(p, stripped)
        i += 1

    if rich:
        enable_field_update(doc)

    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)


def main():
    count = 0
    for d in SRC_DIRS:
        src_dir = os.path.join(ROOT, d)
        if not os.path.isdir(src_dir):
            continue
        for name in sorted(os.listdir(src_dir)):
            if not name.endswith(".md") or name in SKIP_FILES:
                continue
            src = os.path.join(src_dir, name)
            dst = os.path.join(OUT_ROOT, d, name[:-3] + ".docx")
            render(src, dst)
            count += 1
            print("  +", os.path.relpath(dst, ROOT))
    print(f"Готово: собрано {count} .docx в _build/docx/")


if __name__ == "__main__":
    main()
